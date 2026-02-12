import io
from typing import Dict, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ReportGeneratorDocx:
    def __init__(self, report_data: Dict[str, Any]):
        self.report_data = report_data
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def generate(self) -> io.BytesIO:
        self._add_header()
        
        if "results" in self.report_data:
            # We follow the same order as frontend if possible, or iterate
            # For simplicity, we iterate known keys or all keys
            results = self.report_data["results"]
            
            # 1. Compliance Sections (Everything except Tables)
            for section_key, section_data in results.items():
                if section_key in ["nutrition_facts", "sweeteners", "additives"]:
                    continue # Handle tables later
                
                self._add_compliance_section(section_key, section_data)

            # 2. Nutrition Facts
            if "nutrition_facts" in results:
                self._add_nutrition_table(results["nutrition_facts"])

            # 3. Detection Tables
            for table_key in ["sweeteners", "additives"]:
                if table_key in results:
                    self._add_detection_table(table_key, results[table_key])

        # Save to stream
        file_stream = io.BytesIO()
        self.doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    def _add_header(self):
        head = self.doc.add_heading('Compliance Report', 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Job ID: {self.report_data.get('job_id', 'N/A')}\n").bold = True
        p.add_run(f"Generated: {self.report_data.get('created_at', 'Now')}")
        
        self.doc.add_page_break()

    def _add_compliance_section(self, title: str, data: Any):
        if not isinstance(data, dict): return
        
        check_results = data.get("check_results") or data.get("results")
        if not check_results or not isinstance(check_results, list): return

        self.doc.add_heading(title.replace("_", " ").title(), level=1)

        for check in check_results:
            p = self.doc.add_paragraph()
            p.space_after = Pt(12)
            
            # Question
            runner = p.add_run(f"Q: {check.get('question')}\n")
            runner.bold = True
            
            # Result
            result = check.get('result', 'UNKNOWN').upper()
            runner = p.add_run(f"Result: {result}\n")
            if result == 'FAIL':
                runner.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            elif result == 'PASS':
                runner.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            else:
                runner.font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)

            # Rationale
            p.add_run(f"Rationale: {check.get('rationale', '')}\n")
            
            # User Comment
            if check.get('user_comment'):
                runner = p.add_run(f"Comment: {check.get('user_comment')}")
                runner.italics = True
                runner.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)

    def _add_nutrition_table(self, data: Any):
        self.doc.add_heading("Nutrition Facts Audit", level=1)

        if not isinstance(data, dict):
            return

        # Nutrient Audits Table
        nutrient_audits = data.get("nutrient_audits", [])
        if nutrient_audits:
            self.doc.add_heading("Nutrient Rounding Audit", level=2)
            table = self.doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            headers = ["Nutrient", "Original", "Expected", "Status", "Message"]
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                # Bold header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for audit in nutrient_audits:
                row = table.add_row()
                row.cells[0].text = audit.get("nutrient_name", "")
                row.cells[1].text = f"{audit.get('original_value', '')} {audit.get('unit', '')}"
                exp = audit.get("expected_value")
                row.cells[2].text = f"{exp} {audit.get('unit', '')}" if exp is not None else "N/A"
                row.cells[3].text = audit.get("status", "")
                row.cells[4].text = audit.get("message", "")

        # Cross-field Audits
        cross_audits = data.get("cross_field_audits", [])
        if cross_audits:
            self.doc.add_heading("Cross-Field Validation", level=2)
            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            headers = ["Check", "Status", "Message"]
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for audit in cross_audits:
                row = table.add_row()
                row.cells[0].text = audit.get("check_name", "")
                row.cells[1].text = audit.get("status", "")
                row.cells[2].text = audit.get("message", "")

    def _add_detection_table(self, title: str, data: Any):
        self.doc.add_heading(title.replace("_", " ").title(), level=1)

        if not isinstance(data, dict):
            return

        detected = data.get("detected", [])
        if not detected:
            self.doc.add_paragraph("No items detected.")
            return

        # Check if any item has quantity field
        has_quantity = any(d.get("quantity") for d in detected)
        cols = 3 + (1 if has_quantity else 0)
        table = self.doc.add_table(rows=1, cols=cols)
        table.style = 'Table Grid'

        headers = ["Name", "Category", "Source"]
        if has_quantity:
            headers.append("Quantity")

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for item in detected:
            row = table.add_row()
            row.cells[0].text = item.get("name", "")
            row.cells[1].text = item.get("category", "")
            row.cells[2].text = item.get("source", "")
            if has_quantity:
                row.cells[3].text = item.get("quantity", "") or ""
